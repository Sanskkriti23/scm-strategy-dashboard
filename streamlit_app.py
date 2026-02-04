
"""
Streamlit version of the SCM Strategy dashboard (now more actionable with insights based on market gaps, feature differentiation, and revenue projections).

Local run:
  pip install -r requirements.txt
  streamlit run streamlit_app.py
"""

from __future__ import annotations
from pathlib import Path
import re
import pandas as pd
from docx import Document
import streamlit as st
import plotly.express as px

DOCX_PATH = Path(__file__).with_name("refer.docx")

# ---------- Parsing helpers ----------
@st.cache_data(show_spinner=False)
def tables_from_docx(path: Path) -> list[pd.DataFrame]:
    doc = Document(str(path))
    dfs: list[pd.DataFrame] = []
    for t in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        if not rows or len(rows) < 2:
            continue
        header = rows[0]
        data = rows[1:]
        dfs.append(pd.DataFrame(data, columns=header))
    return dfs


@st.cache_data(show_spinner=False)
def bullets_from_docx(path: Path, section_start: str) -> list[str]:
    doc = Document(str(path))
    bullets: list[str] = []
    in_section = False
    section_re = re.compile(r"^\s*Section\s+[A-Z]\b", re.I)

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue

        if txt.lower().startswith(section_start.lower()):
            in_section = True
            continue

        if in_section and section_re.match(txt):
            break

        if in_section and (txt.startswith("•") or txt.startswith("-") or txt.startswith("–")):
            bullets.append(txt.lstrip("•-–").strip())

    return bullets


def qual_to_score(s: str) -> int:
    s = (s or "").lower()
    if any(k in s for k in ["very complex", "enterprise only", "high tco", "expensive", "heavy infra"]):
        return 5
    if any(k in s for k in ["complex", "weak", "limited", "vertical only", "no execution", "no native"]):
        return 3
    if any(k in s for k in ["basic", "add-ons", "community"]):
        return 2
    return 4


dfs = tables_from_docx(DOCX_PATH)
if len(dfs) < 6:
    st.error("Expected 6 tables in the DOCX (Competitors, Gap Grid, AI Map, Segments, Companies, Revenue).")
    st.stop()

df_comp, df_gap, df_ai, df_segments, df_companies, df_revenue = dfs[:6]

insights_d = bullets_from_docx(DOCX_PATH, "Section D")
pilot_playbook = bullets_from_docx(DOCX_PATH, "E4. Pilot Playbook")

DEFAULT_PLAYBOOK = [
    "Week 1–2: Baseline mapping (current tools, SKU count, warehouses, order volume).",
    "Week 3–4: Inventory + PO/GRN live; document AI ingestion; minimum integrations.",
    "Week 5–8: Order + dispatch tracking; control tower; alerting.",
    "Week 9–12: Forecasting + reorder suggestions; KPI review; case study sign-off.",
]

df_comp_vis = df_comp.copy()
df_comp_vis["Complexity/Cost (proxy)"] = df_comp_vis.get("Key Limitation", pd.Series([""]*len(df_comp_vis))).apply(qual_to_score)
df_comp_vis["AI Strength (proxy)"] = df_comp_vis.get("AI Strength", pd.Series([""]*len(df_comp_vis))).apply(qual_to_score)

df_ai_vis = df_ai.copy()
df_ai_vis["SME Differentiation (proxy)"] = df_ai_vis.get("SME Differentiation", pd.Series([""]*len(df_ai_vis))).apply(qual_to_score)

# ---------- UI ----------
st.set_page_config(page_title="SCM Strategy Dashboard", layout="wide")
st.title("Supply Chain ERP Strategy – Interactive Dashboard")
st.caption(f"Data source: {DOCX_PATH.name}")

tabs = st.tabs(["Competitors", "Market Gaps → Features", "AI Feature Map", "Bangalore GTM", "Revenue Model", "Strategy Insights"])

# --- Competitors ---
with tabs[0]:
    c1, c2 = st.columns([1, 2])
    segments = ["All"] + sorted(df_comp.get("Primary Segment", pd.Series(dtype=str)).dropna().unique().tolist())
    with c1:
        seg = st.selectbox("Filter by primary segment", segments, index=0)
    with c2:
        q = st.text_input("Search (vendor / limitation / opportunity)", "")

    dff = df_comp_vis.copy()
    if seg != "All" and "Primary Segment" in dff.columns:
        dff = dff[dff["Primary Segment"] == seg]

    if q:
        s = q.lower()
        mask = False
        for col in [c for c in ["Vendor", "Key Limitation", "Strategic Opportunity", "Coverage Depth"] if c in dff.columns]:
            mask = mask | dff[col].astype(str).str.lower().str.contains(s, na=False)
        dff = dff[mask]

    st.subheader("Competitor Landscape Matrix")
    st.dataframe(dff[df_comp.columns], use_container_width=True, height=360)

    st.subheader("AI vs Complexity (text-proxy)")
    if "Vendor" in dff.columns:
        fig = px.scatter(dff, x="Complexity/Cost (proxy)", y="AI Strength (proxy)", hover_name="Vendor")
        st.plotly_chart(fig, use_container_width=True)

    # Feature Prioritization Bubble Chart (Handling Non-Numeric)
    st.subheader("Feature Prioritization vs Feasibility & Market Demand")
    dff['Strategic Opportunity'] = pd.to_numeric(dff.get('Strategic Opportunity', pd.Series([0] * len(dff))), errors='coerce').fillna(0)
    fig = px.scatter(
        dff,
        x="Complexity/Cost (proxy)",
        y="AI Strength (proxy)",
        size="Strategic Opportunity",  # bubble size represents impact
        color="Coverage Depth",  # use color to highlight strategic opportunities
        hover_name="Vendor",
        title="Features by Feasibility, Market Demand & Impact",
    )
    st.plotly_chart(fig, use_container_width=True)

# --- Market gaps ---
with tabs[1]:
    st.subheader("Market Gap → Product Feature → Business Value Grid")
    st.dataframe(df_gap, use_container_width=True, height=360)

    q = st.text_input("Search within gaps/features/value", "", key="gap_search")
    dff = df_gap.copy()
    if q:
        s = q.lower()
        mask = False
        for col in [c for c in ["Market Gap", "Product Opportunity", "Business Value"] if c in dff.columns]:
            mask = mask | dff[col].astype(str).str.lower().str.contains(s, na=False)
        dff = dff[mask]
    st.subheader("Filtered view")
    st.dataframe(dff, use_container_width=True, height=360)

# --- AI map ---
with tabs[2]:
    features = sorted(df_ai.get("AI Feature", pd.Series(dtype=str)).dropna().unique().tolist())
    c1, c2 = st.columns([1, 2])
    with c1:
        selected = st.multiselect("Select AI features", features, default=features[:6])
        min_diff = st.slider("Minimum SME differentiation (proxy)", 1, 5, 2, 1)

    dff = df_ai_vis.copy()
    if selected and "AI Feature" in dff.columns:
        dff = dff[dff["AI Feature"].isin(selected)]
    dff = dff[dff["SME Differentiation (proxy)"] >= min_diff]

    with c2:
        st.subheader("AI Features Table")
        st.dataframe(dff[df_ai.columns], use_container_width=True, height=360)

    if "AI Feature" in dff.columns:
        st.subheader("SME Differentiation (proxy)")
        fig = px.bar(dff.sort_values("SME Differentiation (proxy)", ascending=False), x="AI Feature", y="SME Differentiation (proxy)")
        st.plotly_chart(fig, use_container_width=True)

# --- GTM ---
with tabs[3]:
    st.subheader("Beachhead Segments (Bangalore)")
    st.dataframe(df_segments, use_container_width=True, height=280)

    st.subheader("Target Company Examples")
    st.dataframe(df_companies, use_container_width=True, height=280)

    st.subheader("90-Day Pilot Playbook")
    for b in (pilot_playbook or DEFAULT_PLAYBOOK):
        st.write(f"• {b}")

# --- Revenue ---
with tabs[4]:
    st.subheader("Revenue Model Mapping")
    st.dataframe(df_revenue, use_container_width=True, height=320)

    st.subheader("Simple MRR Estimator")
    r1, r2, r3 = st.columns(3)
    with r1:
        modules = st.multiselect("Selected modules", ["Inventory", "Orders", "Transport", "CRM", "Dashboards", "S&OP"], default=["Inventory", "Orders"])
        module_price = st.number_input("Module price (per module / month)", min_value=0.0, value=150.0, step=10.0)
    with r2:
        seats = st.number_input("Active users (seat-based)", min_value=0, value=10, step=1)
        seat_price = st.number_input("Seat price (per user / month)", min_value=0.0, value=12.0, step=1.0)
    with r3:
        ai_units = st.number_input("AI usage units (per month)", min_value=0, value=0, step=100)
        ai_price = st.number_input("AI unit price", min_value=0.0, value=0.02, step=0.01, format="%.2f")

    module_rev = len(modules) * module_price
    seat_rev = seats * seat_price
    ai_rev = ai_units * ai_price
    total = module_rev + seat_rev + ai_rev
    st.metric("Estimated MRR", f"{total:,.2f}", help=f"modules={module_rev:,.2f}, seats={seat_rev:,.2f}, AI={ai_rev:,.2f}")

    if "Revenue Stream" in df_revenue.columns:
        tmp = df_revenue.copy()
        tmp["count"] = 1
        st.subheader("Expansion triggers (counts by revenue stream)")
        fig = px.bar(tmp, x="Revenue Stream", y="count")
        st.plotly_chart(fig, use_container_width=True)

# --- Insights ---
with tabs[5]:
    st.subheader("Strategy Validation Insights (Section D)")
    if insights_d:
        for b in insights_d:
            st.write(f"• {b}")
    else:
        st.info("No bullets extracted from Section D in the DOCX.")
